import os
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor


TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?(?:\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")


def add_inline_text(paragraph, text):
    parts = re.split(r"(\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|_[^_]+_)", text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        if (part.startswith("**") and part.endswith("**")) or (part.startswith("__") and part.endswith("__")):
            run.bold = True
            run.text = part[2:-2]
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):
            run.italic = True
            run.text = part[1:-1]
        else:
            run.text = part


def split_table_row(line):
    stripped = line.strip().strip("|")
    if not stripped:
        return []
    return [cell.strip() for cell in stripped.split("|")]


def is_table_start(lines, index):
    if index + 1 >= len(lines):
        return False
    current = lines[index].strip()
    next_line = lines[index + 1].strip()
    return "|" in current and TABLE_SEPARATOR_RE.match(next_line) is not None


def add_table(document, rows):
    if not rows:
        return

    column_count = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"

    for row_index, row in enumerate(rows):
        for column_index in range(column_count):
            cell = table.cell(row_index, column_index)
            cell.text = row[column_index] if column_index < len(row) else ""
            if row_index == 0:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True


def add_code_block(document, code_lines):
    for line in code_lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)


def convert_markdown_to_docx(input_md):
    if not os.path.exists(input_md):
        print(f"Error: {input_md} not found.")
        return False

    output_docx = os.path.splitext(input_md)[0] + ".docx"
    print(f"Processing: {os.path.basename(input_md)}...")

    try:
        with open(input_md, "r", encoding="utf-8") as file_handle:
            lines = file_handle.read().splitlines()

        document = Document()

        # Optimize document-wide whitespace
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(10)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0

        section = document.sections[0]
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Pt(36)
        section.bottom_margin = Pt(36)
        section.left_margin = Pt(36)
        section.right_margin = Pt(36)

        in_code_block = False
        code_lines = []
        pending_blank = False
        index = 0

        while index < len(lines):
            line = lines[index]
            stripped = line.strip()

            if stripped.startswith("```"):
                if in_code_block:
                    add_code_block(document, code_lines)
                    code_lines = []
                    in_code_block = False
                else:
                    in_code_block = True
                index += 1
                continue

            if in_code_block:
                code_lines.append(line)
                index += 1
                continue

            if not stripped:
                pending_blank = True
                index += 1
                continue

            if pending_blank:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                pending_blank = False

            heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                paragraph = document.add_paragraph(style=f"Heading {level}")
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                add_inline_text(paragraph, heading_match.group(2))
                index += 1
                continue

            if is_table_start(lines, index):
                table_rows = [split_table_row(lines[index])]
                index += 2
                while index < len(lines):
                    candidate = lines[index].strip()
                    if not candidate or "|" not in candidate:
                        break
                    if candidate.startswith("```"):
                        break
                    table_rows.append(split_table_row(lines[index]))
                    index += 1
                add_table(document, table_rows)
                continue

            bullet_match = re.match(r"^\s*[-*+]\s+(.*)$", line)
            if bullet_match:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                add_inline_text(paragraph, bullet_match.group(1))
                index += 1
                continue

            number_match = re.match(r"^\s*\d+\.\s+(.*)$", line)
            if number_match:
                paragraph = document.add_paragraph(style="List Number")
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                add_inline_text(paragraph, number_match.group(1))
                index += 1
                continue

            if stripped in {"---", "***", "___"}:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                index += 1
                continue

            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline_text(paragraph, line)
            index += 1

        if in_code_block and code_lines:
            add_code_block(document, code_lines)

        document.save(output_docx)
        print(f"Done: {output_docx}")
        return True

    except Exception as exception:
        print(f"Failed to convert {input_md}: {exception}")
        return False


if __name__ == "__main__":
    paths = [
        r"d:\KrasneLudy\DOKUMENTACJA.md",
        r"d:\KrasneLudy\POPRAWNOŚĆ_ROZWIĄZANIA.md",
    ]
    for path in paths:
        convert_markdown_to_docx(path)
